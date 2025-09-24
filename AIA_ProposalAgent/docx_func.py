from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import json
import os

# Define directories for images and data
IMAGES_DIR = os.path.join(os.path.dirname(__file__), "images")
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")

if os.path.exists("test_doc.docx"):
  os.remove("test_doc.docx")

def add_title(doc, text, size=48, font_name='Calibri Bold', alignment=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    font = run.font
    font.size = Pt(size)
    font.name = font_name
    return p

def add_heading(doc, text, size=16, font_name='Calibri Bold', alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    font = run.font
    font.size = Pt(size)
    font.name = font_name
    return p

def add_subheading(doc, text, size=14, font_name='Calibri Bold', alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    font = run.font
    font.size = Pt(size)
    font.name = font_name
    return p

def add_body_text(doc, text, size=11, font_name='Calibri', alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=None, spacing_before=None, spacing_after=None):
    if not text:
        return doc.add_paragraph()
        
    paragraphs = text.split('\n')
    first_p = None
    
    for p_text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = alignment
        
        # Check if line starts with a dash
        if p_text.lstrip().startswith('- '):
            # Convert to bullet point
            p.style = 'ListBullet'
            # Remove the dash
            p_text = p_text.lstrip('- ')
            # Handle indentation for sub-bullets
            leading_spaces = len(p_text) - len(p_text.lstrip())
            if leading_spaces > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (leading_spaces // 2))
        
        run = p.add_run(p_text.lstrip())
        font = run.font
        font.size = Pt(size)
        font.name = font_name
        
        if line_spacing:
            p.paragraph_format.line_spacing_rule = line_spacing
        if spacing_before is not None:
            p.paragraph_format.space_before = Pt(spacing_before)
        if spacing_after is not None:
            p.paragraph_format.space_after = Pt(spacing_after)
            
        if first_p is None:
            first_p = p
            
    return first_p

def add_bullet_point(doc, text, size=11, font_name='Calibri'):
    p = doc.add_paragraph(text, style='ListBullet')
    if p.runs:
        font = p.runs[0].font
        font.size = Pt(size)
        font.name = font_name
    return p

def add_bullet_points_from_list(doc, items_list, size=11, font_name='Calibri'):
    if isinstance(items_list, list):
        for item in items_list:
            if item != "Not specified":
                add_bullet_point(doc, item, size, font_name)
    else:
        add_body_text(doc, "No items found")

# The initial information table on the first page
def create_info_table(doc, data_dict):
    # Create table with the number of rows matching the data dictionary
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Fill table with data
    for i, (key, value) in enumerate(data_dict.items()):
        # Set header cell
        header_cell = table.rows[i].cells[0]
        header_cell.width = Inches(2.5)
        
        # Clear any existing content
        for paragraph in header_cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        
        # Add text with direct formatting
        run = header_cell.paragraphs[0].add_run(key)
        font = run.font
        font.size = Pt(11)
        font.name = 'Calibri'
        font.bold = True
        font.color.rgb = RGBColor(255, 255, 255)  # White text
        
        # Set value cell with proper formatting
        value_cell = table.rows[i].cells[1]
        
        # Clear any existing content
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        
        # Add text with direct formatting
        run = value_cell.paragraphs[0].add_run(str(value) if value is not None else "")
        font = run.font
        font.size = Pt(11)
        font.bold = False
        font.name = 'Calibri'  # Set font to Calibri
        
        # Add shading to header cells
        shading_elm = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
        header_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add alternating shading to value cells
        shading_color = "cccccc" if i % 2 == 0 else "d9d9d9"
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shading_color))
        value_cell._tc.get_or_add_tcPr().append(shading_elm)
    
    return table

# Creates the invisible header table containing logo and company name
def create_header_table(doc, logo_path=None, company_name=None, company_name_font='Calibri'):
    table = doc.add_table(rows=1, cols=2)
    
    # Add logo if path is provided
    if logo_path:
        pic_cell = table.rows[0].cells[0].paragraphs[0]
        run = pic_cell.add_run()
        run.add_picture(logo_path, width=Inches(2))
    
    # Add company name if provided
    if company_name:
        text_cell = table.rows[0].cells[1]
        # Clear any existing text
        for paragraph in text_cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        
        # Add company name with direct formatting
        run = text_cell.paragraphs[0].add_run(company_name)
        font = run.font
        font.size = Pt(28)
        font.name = 'Calibri'
        font.bold = 'Bold' in company_name_font
    
    return table

# Creates a general table other functions can use
def create_general_table(doc, headers, data_rows):
    # Determine how many data rows we need to create
    actual_data_rows = len(data_rows)
    total_data_rows = max(2, actual_data_rows)  # At least 2 data rows
    
    # Create a table with header row + required data rows
    table = doc.add_table(rows=1 + total_data_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Set header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        
        # Add dark gray background
        shading_elm = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Add data rows
    for row_idx in range(total_data_rows):
        shading_color = "d9d9d9"  # Same color for all data rows
        
        for col_idx in range(len(headers)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            
            # Add data if available, otherwise leave cell empty but properly formatted
            if row_idx < actual_data_rows and col_idx < len(data_rows[row_idx]):
                value = data_rows[row_idx][col_idx]
                run = cell.paragraphs[0].add_run(str(value))
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            else:
                # Empty cell, but ensure it has a paragraph for consistent formatting
                run = cell.paragraphs[0].add_run("")
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            
            # Apply shading regardless of content
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shading_color))
            cell._tc.get_or_add_tcPr().append(shading_elm)
    return table

def create_change_log_table(doc, author):
    current_date = datetime.now().strftime("%d/%m/%Y")
    headers = ["Revision", "Change Description", "Approval Date", "Author"]
    data_rows = [["1.0", "Initial Draft", current_date, author]]
    return create_general_table(doc, headers, data_rows)

def create_timeline_table(doc, timeline_data):
    # Create headers
    headers = ["Milestone", "Description", "Estimated Time (Days)"]
    data_rows = []
    total_duration = timeline_data.get("TOTAL_DURATION", "")

    # Process timeline data
    if timeline_data and isinstance(timeline_data, dict):
        # Extract milestone information from the MILESTONES list in the data
            for i, milestone in enumerate(timeline_data["MILESTONES"], 1):
                if isinstance(milestone, dict):
                    # Use ordered number as the milestone number
                    milestone_number = str(i)
                    description = milestone.get("DESCRIPTION", "")
                    
                    # Extract estimated time (strip any "days" text if present)
                    estimated_time = milestone.get("ESTIMATED_TIME", milestone.get("ESTIMATED_TIME", ""))
                    if isinstance(estimated_time, str):
                        estimated_time = estimated_time.replace("days", "").replace("day", "").strip()
                    
                    data_rows.append([milestone_number, description, estimated_time])
    
    # If no valid data rows were created, add placeholders
    if not data_rows:
        data_rows = [
                ["1", "Development Phase", ""], 
                ["2", "Implementation Phase", ""], 
                ["3", "Support & Maintenance", ""]
            ]
    
    create_general_table(doc, headers, data_rows)
    add_body_text(doc, f"\nTotal Duration: {total_duration}")

def create_budget_table(doc, budget_data, total_duration):
    if isinstance(total_duration, str):
        total_duration = total_duration.replace("days", "").replace("day", "").strip()
    try:
        time = int(str(total_duration).strip()) if total_duration else 1
    except (ValueError, TypeError):
        time = 1

    day_rate = 1600
    cost = time * day_rate
    headers = ["Category", "Time (Days)", "Day Rate", "Cost ($)"]
    data_rows = [["Developer Effort", str(time), str(day_rate), str(cost)]]
    additional_costs = budget_data.get("ADDITIONAL_COSTS", [])
    total_cost = cost
     
    # Add additional costs if any
    if additional_costs and isinstance(additional_costs, list):
        for cost_item in additional_costs:
            if isinstance(cost_item, dict):
                category = cost_item.get("CATEGORY", "")
                amount = cost_item.get("AMOUNT", "")
                if category and amount:
                    try:
                        amount = float(amount)
                        total_cost += amount
                        data_rows.append([category, "", "", str(amount)])
                    except ValueError:
                        continue

    # Add total and GST rows
    data_rows.append(["", "", "Total Cost", str(total_cost)])
    data_rows.append(["", "", "+ 10% GST", str(round(total_cost * 1.1, 2))])
    create_general_table(doc, headers, data_rows)

def add_delivery_team_details(doc, delivery_team_text):
    # Load team member database from JSON
    try:
        json_path = os.path.join(DATA_DIR, 'deliveryteam.json')
        with open(json_path, 'r') as file:
            data = json.load(file)
            team_members_database = data["TEAM_MEMBERS"]
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"Error loading deliveryteam.json: {str(e)}")
        team_members_database = {}
    
    # Define name variations and map them to full names
    name_variations = {
        "Sam": "Samuel Cunningham",
        "Sam Cunningham:": "Samuel Cunningham",
        "Samuel": "Samuel Cunningham",
        "Sean": "Sean Oldenburger",
        "Lindsey": "Lindsey Hershman"
    }
    
    # Look for names in the document text
    found_names = []
    
    if isinstance(delivery_team_text, dict) and "TEAM_MEMBERS" in delivery_team_text:
        for member in delivery_team_text.get("TEAM_MEMBERS", []):
            member_name = member.get("NAME", "")
            if member_name:
                    found_names.append(member_name)
    
    # Process found names to get full names
    full_names = []
    for name in found_names:
        full_name = name_variations.get(name, name)
        if full_name not in full_names:
            full_names.append(full_name)
    
    # Display team members
    for full_name in full_names:
        if full_name in team_members_database:
            member_info = team_members_database[full_name]
            
            # Create a table for the member's image and info (1 row, 2 columns)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths
            table.columns[0].width = Inches(1.5)  # Width for image column
            table.columns[1].width = Inches(5.5)  # Width for text column
            
            # First cell for image
            img_cell = table.cell(0, 0)
            img_paragraph = img_cell.paragraphs[0]
            
            # Add the image
            try:
                image_name = member_info.get("IMAGE", "")
                image_path = os.path.join(IMAGES_DIR, image_name)
                if os.path.exists(image_path):
                    img_paragraph.add_run().add_picture(image_path, width=Inches(1.3))
                    print(f"Added team member image: {image_path}")
                else:
                    print(f"Team member image not found: {image_path}")
                    img_paragraph.add_run(f"[Image not found: {full_name}]")
            except Exception:
                print(f"Error adding team member image: {str(e)}")
                img_paragraph.add_run(f"[Image error: {full_name}]")
            
            # Second cell for text - clear any existing paragraphs
            for paragraph in table.cell(0, 1).paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            
            # Add description with proper formatting
            text_paragraph = table.cell(0, 1).paragraphs[0]
            text_paragraph.text = member_info["DESCRIPTION"]
            
            # Apply consistent formatting
            for run in text_paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            
            # Add name below the table
            name_para = add_body_text(doc, full_name, font_name='Calibri Bold', size=11)
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add space after each member
            doc.add_paragraph("")
        else:
            # Fallback for unknown team members
            add_subheading(doc, full_name, size=12)
            add_body_text(doc, "No additional information available.")
            doc.add_paragraph("")

def add_past_projects_section(doc, past_projects_text):    
    # Load project descriptions from JSON file
    try:
        json_path = os.path.join(DATA_DIR, 'pastprojects.json')
        with open(json_path, 'r') as file:
            project_data = json.load(file)
            project_descriptions = {k: v["DESCRIPTION"] for k, v in project_data["PROJECTS"].items()}
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"Error loading pastprojects.json: {str(e)}")
        project_descriptions = {}
    
    # Display past projects with descriptions
    if isinstance(past_projects_text, dict) and "PAST_PROJECTS" in past_projects_text:
        past_projects = past_projects_text["PAST_PROJECTS"]
        if past_projects and len(past_projects) > 0:
            for project in past_projects:
                project_name = project.get("PROJECT_NAME", "")
                matching_key = None
                description = ""
                for key in project_descriptions:
                    if (key in project_name or project_name in key) and matching_key is None:
                        matching_key = key
                        description = project_descriptions[key]
                if matching_key:
                    add_bullet_point(doc, project_name, size=11, font_name='Calibri Bold')
                    p = doc.add_paragraph()
                    p.style = 'ListBullet'
                    p.paragraph_format.left_indent = Inches(0.5)
                    run = p.add_run(description)
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
        else:
            add_body_text(doc, "No similar past projects were identified for the current requirements.", size=11)
    else:
        add_body_text(doc, "Past project information is not available.", size=11)

def add_page_break(doc):
    doc.add_page_break()
    return doc